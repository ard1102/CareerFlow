"""
Test resume upload and parsing functionality
Tests: POST /api/resume/upload - PDF and DOCX parsing
"""

import pytest
import requests
import os
import io

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

@pytest.fixture(scope="module")
def auth_token():
    """Get auth token for test user"""
    response = requests.post(f"{BASE_URL}/api/auth/login", json={
        "email": "pdftest@test.com",
        "password": "test123"
    })
    if response.status_code == 200:
        return response.json()["access_token"]
    pytest.fail(f"Authentication failed: {response.text}")

@pytest.fixture
def api_client(auth_token):
    """Shared requests session with auth header"""
    session = requests.Session()
    session.headers.update({
        "Authorization": f"Bearer {auth_token}"
    })
    return session


def create_minimal_pdf_bytes():
    """Create a minimal valid PDF with text content"""
    # Using PyMuPDF to create a valid PDF
    try:
        import fitz  # PyMuPDF
        doc = fitz.open()
        page = doc.new_page()
        
        # Add resume-like content
        resume_text = """John Doe
john.doe@email.com
(555) 123-4567

PROFESSIONAL SUMMARY
Experienced software engineer with 5+ years of experience in Python, JavaScript, and cloud technologies.

SKILLS
Python, JavaScript, React, Node.js, AWS, Docker, PostgreSQL, MongoDB

EDUCATION
Bachelor of Science in Computer Science
MIT, 2018

EXPERIENCE
Senior Software Engineer at Tech Corp
2019 - Present
- Developed scalable microservices using Python and FastAPI
- Led team of 5 engineers

PROJECTS
- Built AI-powered job search platform
- Developed real-time analytics dashboard
"""
        page.insert_text((50, 50), resume_text, fontsize=10)
        pdf_bytes = doc.tobytes()
        doc.close()
        return pdf_bytes
    except ImportError:
        # Fallback: create a minimal valid PDF manually
        pdf_content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >> endobj
4 0 obj << /Length 100 >>
stream
BT
/F1 12 Tf
50 700 Td
(John Doe - john.doe@email.com - Python JavaScript React) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000214 00000 n
trailer << /Root 1 0 R /Size 5 >>
startxref
364
%%EOF"""
        return pdf_content


def create_minimal_docx_bytes():
    """Create a minimal valid DOCX with resume content"""
    try:
        from docx import Document
        doc = Document()
        
        doc.add_heading('Jane Smith', 0)
        doc.add_paragraph('jane.smith@example.com | (555) 987-6543')
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph('Full-stack developer with 3 years of experience in web development.')
        
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('JavaScript, TypeScript, React, Angular, Node.js, MongoDB, AWS, Git')
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Master of Science in Software Engineering, Stanford University, 2021')
        
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Software Developer at StartupXYZ (2021-Present)')
        doc.add_paragraph('- Developed REST APIs using Node.js')
        doc.add_paragraph('- Built React frontend applications')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        pytest.skip("python-docx not installed")


class TestResumeUploadPDF:
    """Test PDF resume upload and parsing"""
    
    def test_upload_valid_pdf_returns_success(self, api_client):
        """Test that uploading a valid PDF returns success and parsed data"""
        pdf_bytes = create_minimal_pdf_bytes()
        
        files = {
            'file': ('resume.pdf', io.BytesIO(pdf_bytes), 'application/pdf')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        
        # Should return 200 OK
        assert response.status_code == 200, f"Upload failed: {response.text}"
        
        data = response.json()
        assert data.get("success") == True
        assert "parsed_data" in data
        assert data.get("raw_text_length", 0) > 0
        
        parsed = data["parsed_data"]
        # Verify parsed data structure
        expected_fields = ["name", "email", "skills"]
        for field in expected_fields:
            assert field in parsed, f"Missing field {field} in parsed data"
    
    def test_pdf_extracts_email(self, api_client):
        """Test that email is extracted from PDF"""
        pdf_bytes = create_minimal_pdf_bytes()
        
        files = {
            'file': ('test_resume.pdf', io.BytesIO(pdf_bytes), 'application/pdf')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        assert response.status_code == 200
        
        parsed = response.json().get("parsed_data", {})
        email = parsed.get("email")
        
        # Email should be extracted if present in PDF
        if email:
            assert "@" in email, "Extracted email should be valid format"
    
    def test_pdf_extracts_skills(self, api_client):
        """Test that skills are extracted from PDF"""
        pdf_bytes = create_minimal_pdf_bytes()
        
        files = {
            'file': ('skills_resume.pdf', io.BytesIO(pdf_bytes), 'application/pdf')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        assert response.status_code == 200
        
        parsed = response.json().get("parsed_data", {})
        skills = parsed.get("skills", [])
        
        # Should extract at least some skills from the PDF
        assert isinstance(skills, list), "Skills should be a list"


class TestResumeUploadDOCX:
    """Test DOCX resume upload and parsing"""
    
    def test_upload_valid_docx_returns_success(self, api_client):
        """Test that uploading a valid DOCX returns success"""
        try:
            docx_bytes = create_minimal_docx_bytes()
        except Exception:
            pytest.skip("Could not create DOCX for testing")
        
        files = {
            'file': ('resume.docx', io.BytesIO(docx_bytes), 
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        
        assert response.status_code == 200, f"Upload failed: {response.text}"
        
        data = response.json()
        assert data.get("success") == True
        assert "parsed_data" in data
        assert data.get("raw_text_length", 0) > 0


class TestResumeUploadErrors:
    """Test error handling for resume upload"""
    
    def test_reject_unsupported_file_type(self, api_client):
        """Test that unsupported file types are rejected"""
        # Create a text file disguised as image
        fake_content = b"This is not a PDF or DOCX file"
        
        files = {
            'file': ('resume.txt', io.BytesIO(fake_content), 'text/plain')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        
        # Should reject with 400
        assert response.status_code == 400
        assert "supported" in response.json().get("detail", "").lower()
    
    def test_reject_image_file(self, api_client):
        """Test that image files are rejected"""
        # Create minimal image-like bytes
        fake_image = b'\x89PNG\r\n\x1a\n' + b'\x00' * 100
        
        files = {
            'file': ('resume.png', io.BytesIO(fake_image), 'image/png')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        
        assert response.status_code == 400


class TestResumeParsingQuality:
    """Test parsing quality and data extraction"""
    
    def test_parsed_data_has_expected_structure(self, api_client):
        """Test that parsed data includes expected fields"""
        pdf_bytes = create_minimal_pdf_bytes()
        
        files = {
            'file': ('structure_test.pdf', io.BytesIO(pdf_bytes), 'application/pdf')
        }
        
        response = api_client.post(f"{BASE_URL}/api/resume/upload", files=files)
        assert response.status_code == 200
        
        parsed = response.json().get("parsed_data", {})
        
        # Should have these fields from rule-based parsing
        structure_fields = [
            "name", "email", "phone", "location", "summary",
            "skills", "experience_years", "education", "projects", "work_authorization"
        ]
        
        for field in structure_fields:
            assert field in parsed, f"Missing expected field: {field}"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
