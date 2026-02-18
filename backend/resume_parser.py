"""
Resume parsing utilities
Extract text from PDF/DOCX and parse with AI
Uses PyMuPDF (fitz) as primary PDF parser for better extraction
Falls back to PyPDF2 if PyMuPDF fails
"""

import io
import re
from typing import Dict
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file using multiple methods for robustness"""
    text = ""
    
    # Method 1: PyMuPDF (fitz) - most reliable for complex PDFs
    try:
        import fitz  # PyMuPDF
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text += page.get_text() + "\n"
        pdf_document.close()
        
        if text.strip():
            logger.info(f"PyMuPDF extracted {len(text)} characters")
            return text.strip()
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")
    
    # Method 2: PyPDF2 - fallback
    try:
        from PyPDF2 import PdfReader
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            logger.info(f"PyPDF2 extracted {len(text)} characters")
            return text.strip()
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    # Method 3: pdfplumber - last resort for scanned/OCR PDFs
    try:
        import pdfplumber
        pdf_file = io.BytesIO(file_bytes)
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            logger.info(f"pdfplumber extracted {len(text)} characters")
            return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
    
    logger.error("All PDF extraction methods failed")
    return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file including tables"""
    try:
        from docx import Document
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract tables (resumes often have tabular data)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return ""

def parse_resume_with_ai(resume_text: str) -> Dict:
    """
    Parse resume text using regex and NLP to extract structured data
    This is a rule-based parser. For better results, use AI in the endpoint.
    """
    result = {
        "name": None,
        "email": None,
        "phone": None,
        "location": None,
        "summary": None,
        "skills": [],
        "experience_years": None,
        "education": [],
        "projects": [],
        "work_authorization": None
    }
    
    lines = resume_text.split('\n')
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, resume_text)
    if email_match:
        result["email"] = email_match.group(0)
    
    # Extract phone
    phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phone_match = re.search(phone_pattern, resume_text)
    if phone_match:
        result["phone"] = phone_match.group(0)
    
    # Extract name (usually first non-empty line)
    for line in lines:
        line = line.strip()
        if line and len(line.split()) <= 4 and not any(char.isdigit() for char in line):
            result["name"] = line
            break
    
    # Extract years of experience
    exp_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+experience',
        r'experience:\s*(\d+)\+?\s*(?:years?|yrs?)'
    ]
    for pattern in exp_patterns:
        match = re.search(pattern, resume_text, re.IGNORECASE)
        if match:
            result["experience_years"] = int(match.group(1))
            break
    
    # Extract common skills
    tech_keywords = [
        'Python', 'Java', 'JavaScript', 'TypeScript', 'React', 'Angular', 'Vue',
        'Node.js', 'Django', 'Flask', 'FastAPI', 'Spring', 'AWS', 'Azure', 'GCP',
        'Docker', 'Kubernetes', 'MongoDB', 'PostgreSQL', 'MySQL', 'Redis',
        'Git', 'CI/CD', 'REST API', 'GraphQL', 'Machine Learning', 'AI',
        'TensorFlow', 'PyTorch', 'Pandas', 'NumPy', 'SQL', 'NoSQL',
        'Microservices', 'Agile', 'Scrum', 'HTML', 'CSS', 'C++', 'C#',
        'Go', 'Rust', 'Swift', 'Kotlin', 'Ruby', 'PHP', 'Scala'
    ]
    
    for tech in tech_keywords:
        if re.search(r'\b' + re.escape(tech) + r'\b', resume_text, re.IGNORECASE):
            result["skills"].append(tech)
    
    # Extract education
    edu_patterns = [
        r'(Bachelor.*?(?:\n|$))',
        r'(Master.*?(?:\n|$))',
        r'(PhD.*?(?:\n|$))',
        r'(B\.?S\.?.*?(?:\n|$))',
        r'(M\.?S\.?.*?(?:\n|$))',
        r'(MBA.*?(?:\n|$))'
    ]
    
    for pattern in edu_patterns:
        matches = re.finditer(pattern, resume_text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            edu = match.group(1).strip()
            if edu and len(edu) < 200:  # Reasonable length
                result["education"].append(edu)
    
    # Remove duplicates
    result["skills"] = list(set(result["skills"]))
    result["education"] = list(set(result["education"]))
    
    # Extract summary (look for "Summary", "About", "Profile" sections)
    summary_patterns = [
        r'(?:SUMMARY|PROFESSIONAL SUMMARY|ABOUT|PROFILE|OBJECTIVE)[\s:]+(.+?)(?=\n\n|\n[A-Z]{3,})',
    ]
    
    for pattern in summary_patterns:
        match = re.search(pattern, resume_text, re.IGNORECASE | re.DOTALL)
        if match:
            summary = match.group(1).strip()
            if 50 < len(summary) < 500:  # Reasonable length
                result["summary"] = summary
                break
    
    return result
